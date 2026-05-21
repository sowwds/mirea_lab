from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "report_examples" / "ПР 1.docx"
AUTHORS = ["Ефремов А.И.", "Лазарев Г.С.", "Никитин А.В."]
TEACHER = "Холмогоров В.В."


TABLE_TITLES = {
    ("Параметр", "Значение"): "Основные характеристики набора данных",
    ("Признак", "Описание"): "Описание признаков набора данных",
    ("Компонент", "Назначение", "Источник"): "Компоненты программной реализации",
    ("Источник", "Количество строк"): "Количество строк по источникам",
    ("Признак", "Количество пропусков", "Доля пропусков, %"): "Количество пропусков по признакам",
    ("Показатель", "Значение"): "Результаты предобработки",
    ("Класс", "Precision", "Recall", "F1-score", "Support"): "Метрики логистической регрессии",
    ("target", "Исходные объекты", "Аугментированные объекты", "Итого"): "Результат аугментации данных",
    ("Модель", "Вероятность болезни"): "Пример усреднения вероятностей при soft-voting",
    ("Модель", "CV F1", "Вес"): "Веса моделей по результатам кросс-валидации",
    ("Модель", "Основные параметры", "Зачем так выбрано"): "Используемые модели и параметры",
    ("Модель", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"): "Сравнение качества моделей",
    ("Модель", "Верно: нет болезни", "Ошибочно: есть болезнь", "Ошибочно: нет болезни", "Верно: есть болезнь"): "Матрицы ошибок в числовом виде",
    ("Набор", "Строк", "Признаков", "Класс 0", "Класс 1"): "Структура real train и real test",
    ("Ориентир", "Модель", "F1"): "Пороги сравнения качества",
    ("Генератор", "Размер", "Строк", "Класс 0", "Класс 1", "Positive rate"): "Сгенерированные синтетические наборы",
    ("Генератор", "Размер", "Mean MAE", "Variance MAE", "Binary MAE", "Correlation MAE", "DCR min", "DCR median"): "Статистическая проверка синтетических данных",
    ("Train source", "Лучшая модель", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"): "Лучшие TSTR-результаты",
}

TOC_PAGES = {
    "1": {
        "Введение": 3,
        "1 Теоретическая часть": 5,
        "1.1 Описание набора данных": 5,
        "1.2 Постановка задачи классификации": 6,
        "1.3 Методы предобработки данных": 6,
        "1.4 Аугментация табличных данных": 8,
        "1.5 Методы визуализации и понижения размерности": 9,
        "1.6 Метрики качества классификации": 10,
        "2 Практическая часть": 11,
        "2.1 Основные компоненты программной реализации": 11,
        "2.2 Загрузка и первичный анализ данных": 11,
        "2.3 Реализованная предобработка": 13,
        "2.4 Проверка пригодности данных на модели": 14,
        "2.5 Аугментация данных": 15,
        "2.6 Визуализация данных": 16,
        "Заключение": 21,
        "Список использованных источников": 23,
        "Приложение А": 24,
    },
    "2": {
        "Введение": 3,
        "1 Теоретическая часть": 4,
        "1.1 Описание задачи и входных данных": 4,
        "1.2 Дискриминативная классификация": 4,
        "1.3 Логистическая регрессия": 4,
        "1.4 Случайный лес": 5,
        "1.5 Градиентный бустинг": 7,
        "1.6 Ансамбль голосования": 8,
        "1.7 Метрики качества": 9,
        "2 Практическая часть": 10,
        "2.1 Используемые модели и параметры": 10,
        "2.2 Анализ признаков перед обучением": 10,
        "2.3 Протокол обучения и оценки": 13,
        "2.4 Сравнение моделей": 13,
        "2.5 Матрицы ошибок": 15,
        "Заключение": 17,
        "Список использованных источников": 18,
        "Приложение А": 19,
    },
    "3": {
        "Введение": 3,
        "1 Теоретическая часть": 4,
        "1.1 Постановка задачи генерации табличных данных": 4,
        "1.2 Проверка TSTR": 4,
        "1.3 Gaussian Mixture Model": 5,
        "1.4 Нейронный denoising-autoencoder": 6,
        "1.5 Экспертная база правил": 6,
        "1.6 Метрики качества синтетических данных": 7,
        "2 Практическая часть": 9,
        "2.1 Используемые данные и протокол эксперимента": 9,
        "2.2 Реализованные генераторы": 9,
        "2.3 Статистическая проверка синтетических данных": 10,
        "2.4 Визуальный анализ синтетических данных": 11,
        "2.5 TSTR-сравнение моделей": 12,
        "Заключение": 15,
        "Список использованных источников": 17,
        "Приложение А": 18,
    },
}


def set_run_font(run, *, bold: bool | None = None, size: int = 14, name: str = "Times New Roman") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_base(paragraph, *, align=None, first_line: bool = True) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1.25) if first_line else None
    fmt.page_break_before = False


def clear_paragraph(paragraph, text: str = "", *, align=None) -> None:
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)


def normalize_numbered_heading(text: str) -> str:
    return re.sub(r"^(\d+(?:\.\d+)*)\.\s+", r"\1 ", text)


def normalize_caption_text(text: str) -> str:
    text = re.sub(r"^Рисунок\s+\d+\s*[-—]\s*", "", text.strip())
    return text.replace(" - ", " — ")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    set_run_font(run)


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        add_field(paragraph, "PAGE")


def enable_field_update(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def remove_body_after_title(document: Document, title_end_paragraph_index: int = 18) -> None:
    body = document._body._element
    keep_last = document.paragraphs[title_end_paragraph_index]._p
    children = list(body)
    keep_index = children.index(keep_last)

    for child in children[keep_index + 1 :]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def ensure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    if "Heading 3" not in [s.name for s in styles]:
        h3 = styles.add_style("Heading 3", 1)
    else:
        h3 = styles["Heading 3"]
    h3.base_style = normal
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h3.font.size = Pt(14)
    h3.font.bold = True


def replace_title(document: Document, *, lab_number: str, topic: str) -> None:
    paragraphs = document.paragraphs
    clear_paragraph(paragraphs[3], f"Практическая работа {lab_number}", align=WD_ALIGN_PARAGRAPH.CENTER)
    clear_paragraph(paragraphs[5], f"Тема: {topic}", align=WD_ALIGN_PARAGRAPH.CENTER)
    clear_paragraph(paragraphs[10], f"Выполнили: {AUTHORS[0]},", align=WD_ALIGN_PARAGRAPH.RIGHT)
    clear_paragraph(paragraphs[11], f"{AUTHORS[1]}, {AUTHORS[2]}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    clear_paragraph(paragraphs[12], f"Принял: {TEACHER}", align=WD_ALIGN_PARAGRAPH.RIGHT)


def extract_lab_number(md: str, fallback: str) -> str:
    match = re.search(r"Практическая работа №\s*(\d+)", md)
    return match.group(1) if match else fallback


def extract_topic(md: str) -> str:
    match = re.search(r"\*\*Тема:\*\*\s*(.+)", md)
    return match.group(1).strip() if match else "Практическая работа"


def body_lines(md: str) -> list[str]:
    lines = md.splitlines()
    for index, line in enumerate(lines):
        if line.startswith("## Содержание"):
            return lines[index:]
    return lines


def add_inline(paragraph, text: str) -> None:
    text = text.replace("`", "")
    text = text.replace(" - ", " — ")
    pos = 0
    for match in re.finditer(r"\*\*([^*]+)\*\*", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_normal_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
    add_inline(paragraph, text)


def add_toc_line(document: Document, text: str, page: int | str = "") -> None:
    text = normalize_numbered_heading(text)
    major = (
        text in {"Введение", "Заключение", "Список использованных источников", "Приложение А"}
        or re.match(r"^\d+\s+", text) is not None
    )
    display = text.upper() if major else text
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = paragraph.add_run(display)
    set_run_font(run)
    paragraph.add_run("\t")
    page_run = paragraph.add_run(str(page))
    set_run_font(page_run)


def add_toc_field(document: Document) -> None:
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_field(paragraph, r'TOC \o "1-2" \h \z \u')


def add_heading(document: Document, level: int, text: str) -> None:
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base(paragraph, align=paragraph.alignment, first_line=False)
    paragraph.paragraph_format.space_before = Pt(0 if level == 1 else 24)
    paragraph.paragraph_format.space_after = Pt(12)
    text = normalize_numbered_heading(text)
    display_text = text.upper() if level == 1 else text
    run = paragraph.add_run(display_text)
    set_run_font(run, bold=True, size=18 if level == 1 else 16 if level == 2 else 14)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw.append(lines[index].strip())
        index += 1

    rows = []
    for row_index, line in enumerate(raw):
        cells = [cell.strip().replace("`", "") for cell in line.strip("|").split("|")]
        if row_index == 1 and all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table_caption(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(f"Таблица {number} — {title}")
    set_run_font(run, size=12)
    run.italic = True


def infer_table_title(rows: list[list[str]]) -> str:
    header = tuple(rows[0]) if rows else ()
    return TABLE_TITLES.get(header, "Сводные данные")


def add_table(document: Document, rows: list[list[str]], number: str) -> None:
    if not rows:
        return
    add_table_caption(document, number, infer_table_title(rows))
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Normal Table"

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.clear()
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, bold=True if row_index == 0 else None, size=12)


def add_image(document: Document, md_line: str, md_path: Path, number: str) -> None:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", md_line.strip())
    if not match:
        return
    caption, image_ref = match.groups()
    image_path = (md_path.parent / image_ref).resolve()
    if not image_path.exists():
        add_normal_paragraph(document, f"Рисунок не найден: {image_ref}")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15))

    caption_paragraph = document.add_paragraph()
    set_paragraph_base(caption_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(6)
    caption_paragraph.paragraph_format.line_spacing = 1.0
    add_inline(caption_paragraph, f"Рисунок {number} — {normalize_caption_text(caption)}")
    for run in caption_paragraph.runs:
        set_run_font(run, bold=True, size=12)


def add_formula_block(document: Document, code: list[str]) -> None:
    document.add_paragraph()
    cleaned = [line.rstrip() for line in code]
    if cleaned and cleaned[-1] and cleaned[-1][-1] not in ".,;:":
        cleaned[-1] += "."
    for line in cleaned:
        paragraph = document.add_paragraph()
        set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        run = paragraph.add_run(line)
        set_run_font(run)
    document.add_paragraph()


def add_listing_caption(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    paragraph.paragraph_format.space_before = Pt(6)
    run = paragraph.add_run(f"Листинг {number} — {title}")
    set_run_font(run, size=12)
    run.italic = True


def add_code_block(document: Document, code: list[str], number: str) -> None:
    add_listing_caption(document, number, "Фрагмент программного кода")
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Normal Table"
    cell = table.cell(0, 0)
    set_cell_border(cell)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(code))
    set_run_font(run, size=10, name="Courier New")


def add_plain_code_block(document: Document, code: list[str]) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    run = paragraph.add_run("\n".join(code))
    set_run_font(run, size=10, name="Courier New")


def render_markdown_body(document: Document, md: str, md_path: Path, toc_pages: dict[str, int]) -> None:
    lines = body_lines(md)
    index = 0
    first_heading = True
    toc_mode = False
    in_code = False
    in_code_lang = ""
    code_lines: list[str] = []
    current_section = "0"
    current_appendix = ""
    table_counts: dict[str, int] = {}
    figure_counts: dict[str, int] = {}
    listing_counts: dict[str, int] = {}

    def next_number(counter: dict[str, int]) -> str:
        key = current_appendix or current_section
        counter[key] = counter.get(key, 0) + 1
        return f"{key}.{counter[key]}" if key != "0" else str(counter[key])

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if line.startswith("```"):
            if in_code:
                if in_code_lang == "python":
                    add_code_block(document, code_lines, next_number(listing_counts))
                else:
                    add_formula_block(document, code_lines)
                code_lines = []
                in_code_lang = ""
                in_code = False
            else:
                in_code = True
                in_code_lang = line[3:].strip().split()[0] if line[3:].strip() else ""
            index += 1
            continue

        if in_code:
            code_lines.append(raw_line)
            index += 1
            continue

        if not line:
            index += 1
            continue

        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows, next_number(table_counts))
            continue

        if line.startswith("!["):
            add_image(document, line, md_path, next_number(figure_counts))
            index += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            hashes, text = heading.groups()
            if hashes == "##":
                if not first_heading:
                    document.add_page_break()
                first_heading = False
                normalized = normalize_numbered_heading(text)
                section_match = re.match(r"^(\d+)\s+", normalized)
                if section_match:
                    current_section = section_match.group(1)
                    current_appendix = ""
                elif normalized.upper().startswith("ПРИЛОЖЕНИЕ А"):
                    current_appendix = "А"
                toc_mode = text.strip().lower() == "содержание"
                add_heading(document, 1, text)
            elif hashes == "###":
                toc_mode = False
                add_heading(document, 2, text)
            else:
                toc_mode = False
                add_heading(document, 3, text)
            index += 1
            continue

        if toc_mode:
            normalized_toc_line = normalize_numbered_heading(line)
            add_toc_line(document, normalized_toc_line, toc_pages.get(normalized_toc_line, ""))
            index += 1
            continue

        if line.startswith("- "):
            while index < len(lines) and lines[index].strip().startswith("- "):
                paragraph = document.add_paragraph(style="List Paragraph")
                set_paragraph_base(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)
                paragraph.paragraph_format.left_indent = Cm(2.25)
                paragraph.paragraph_format.first_line_indent = None
                add_inline(paragraph, "– " + lines[index].strip()[2:])
                index += 1
            continue

        paragraph_parts = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if (
                not next_line
                or next_line.startswith(("##", "|", "![", "```", "- "))
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            paragraph_parts.append(next_line)
            index += 1

        add_normal_paragraph(document, " ".join(paragraph_parts))

    if in_code and code_lines:
        if in_code_lang == "python":
            add_code_block(document, code_lines, next_number(listing_counts))
        else:
            add_formula_block(document, code_lines)


def build_report(md_path: Path, output_path: Path, fallback_lab_number: str) -> None:
    md = md_path.read_text(encoding="utf-8")
    document = Document(str(TEMPLATE))
    ensure_styles(document)
    configure_sections(document)
    enable_field_update(document)
    replace_title(
        document,
        lab_number=extract_lab_number(md, fallback_lab_number),
        topic=extract_topic(md),
    )
    remove_body_after_title(document)
    render_markdown_body(document, md, md_path, TOC_PAGES.get(extract_lab_number(md, fallback_lab_number), {}))
    document.save(output_path)


def main() -> int:
    build_report(ROOT / "reports" / "lab1" / "report_lab1.md", ROOT / "reports" / "lab1" / "report_lab1.docx", "1")
    build_report(ROOT / "reports" / "lab2" / "report_lab2.md", ROOT / "reports" / "lab2" / "report_lab2.docx", "2")
    build_report(ROOT / "reports" / "lab3" / "report_lab3.md", ROOT / "reports" / "lab3" / "report_lab3.docx", "3")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
